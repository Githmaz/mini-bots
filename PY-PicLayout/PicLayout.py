import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_centered_image_doc(img1, img2, img3, output_name="three_images.docx"):

    print(f"size = { img1.size}")
    # Ensure /doc folder exists
    output_folder = "doc"
    os.makedirs(output_folder, exist_ok=True)
    
    # Create document
    doc = Document()

    # Set to A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Add and center each image
    for img_path in [img1, img2, img3]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(5))  # Adjust width as needed
        doc.add_paragraph()  # Spacer

    # Save to /doc folder
    output_path = os.path.join(output_folder, output_name)
    doc.save(output_path)
    print(f"✅ Document saved at: {output_path}")

# Example usage
create_centered_image_doc("black.png", "blue.png", "red.png")
